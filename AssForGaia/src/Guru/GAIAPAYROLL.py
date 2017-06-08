#用于模拟浏览器
from selenium import webdriver
import unittest, time#延时用
#用于读取docx文档
import docx
from docutils.nodes import paragraph
#beautiful 用于解析网页
import bs4 as bs
import http.client as httplib
#快速请求
import httplib2
#普通请求
import urllib.request
import http.cookiejar as cookielib
import urllib.parse
import json #json数据格式，用于写入数据库链接配置信息
from types import CodeType

t1 = time.clock()

def requestNeedurl(needurl,postData,headers,CookieOpener):
    needurl = needurl
    headers['Accept'] = 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8' 
    if headers.get('Content-Type') == None:
        pass
    else:
        del headers['Content-Type']
    encodepostData = urllib.parse.urlencode(postData).encode('utf-8')
    req = urllib.request.Request(url = needurl,data = encodepostData,headers = headers)
    response = CookieOpener.open(req)
    return response
def getpermanentID(rooturl,headers,loginurl): 
    req = urllib.request.Request(url = loginurl,headers = headers)
    response = urllib.request.urlopen(req)
    soup = bs.BeautifulSoup(response.read().decode(),"lxml")
    permanentID = soup.form.find(id="permanentId")["value"]
    return permanentID
def loginInSystem(requestInfo):
    permanentID = getpermanentID(rooturl=requestInfo.rooturl,headers=headers,loginurl=requestInfo.loginurl)
    postData = {'account':requestInfo.webname,'password':requestInfo.webpsw,'cultureCode':'zh-CN','permanentId':permanentID}
    print('requestInfo.loginurl:%s'%(requestInfo.loginurl))
    response = requestNeedurl(needurl  = requestInfo.loginurl ,postData = postData,headers = headers,CookieOpener = requestInfo.CookieOpener)
    fullcookies = ''
    for item in cookie:
            singlecookie = item.name + '=' + item.value + ';'
            fullcookies = fullcookies + singlecookie
    requestInfo.headers['Cookie'] = fullcookies        
    return response
def requestNeedurlByJsonPost(needurl,postData,headers,host,port):
    needurl = needurl
    headers['Content-Type'] = 'application/json'
    headers['Accept'] = 'application/json, text/javascript, */*; q=0.01'
    conn = httplib.HTTPConnection(host,port)
    conn.request("POST", needurl, json.JSONEncoder().encode(postData), headers)
    response = conn.getresponse()



class UrlRequest:
    def __init__(self,rooturl,host,port,webname,webpsw,tmpFilePath,showMode,cookie,headers,loginurl,CookieOpener):
        self.rooturl = rooturl
        self.host = host
        self.port = port
        self.webname = webname
        self.webpsw = webpsw
        self.tmpFilePath = tmpFilePath 
        self.showMode = showMode
        self.cookie = cookie
        self.headers = headers
        self.loginurl = loginurl
        self.CookieOpener = CookieOpener        
def checkContentInMimeTypes(checkThing,mimetypepath,addSet):
    ''' for check content like xls in mimetype configfile,and delete old config,
        set new config(can download with no windows),return update result and old  content which was overwrite
        checkThing : for xls、xlsx、pdf,be sure as same as addSet config
        mimetypepath:the config file in local for webdirver firefox
        addset:for checkthings ,you need each set for checktings,dict it is 
    '''
    checkState = 0
    oldCheckContent = ''
    fileobject = open(mimetypepath,'rb')
    fileReadContent = fileobject.readlines()
    fileobject.close()
    
    fileobject = open(mimetypepath,'wb')
    checkName = r'urn:mimetype:application/' + checkThing
    for line in fileReadContent:
        lineContent = line.decode('utf-8')              
        if checkState > 0:
            if  r'</RDF:Description>' in lineContent:
                checkState = 0
                oldCheckContent = oldCheckContent + lineContent
            else:
                checkState += 1
                oldCheckContent = oldCheckContent + lineContent
        else:
            if checkName in lineContent:
                checkState = 1
                oldCheckContent = oldCheckContent + lineContent
                if  r'</RDF:Description>' in lineContent:
                    checkState = 0
                else:
                    pass
            else:
                if r'</RDF:RDF>' in lineContent:
                    
                    fileobject.write(addSet.get(checkThing).encode('utf-8')) 
                    fileobject.write(line)
                else:
                    fileobject.write(line)
    fileobject.close()
        
    fileobject = open(mimetypepath,'rb')
    newReadContent = fileobject.read().decode('utf-8')
    fileobject.close()                
    return  newReadContent,oldCheckContent    
def writeMimetype(mimetypepath): 
    ''' use def checkContentInMimeTypes overwrite  mimetype config file one by one
    mimetypepath:the config file in local for webdirver firefox '''   
    mimetypepath = r'C:\Users\Administrator\AppData\Roaming\Mozilla\Firefox\Profiles\hdjophne.default\mimeTypes.rdf'
    addSet = {} 
    addSet['xls'] =  r'''<RDF:Description RDF:about="urn:mimetype:application/xls" NC:value="application/xls" NC:fileExtensions="xls">  </RDF:Description>
        <RDF:Description RDF:about="urn:mimetype:handler:application/xls" NC:saveToDisk="true" NC:alwaysAsk="false" ></RDF:Description>'''
    addSet['xlsx'] = r'''<RDF:Description RDF:about="urn:mimetype:application/xlsx" NC:value="application/xlsx" NC:fileExtensions="xlsx">  </RDF:Description>
        <RDF:Description RDF:about="urn:mimetype:handler:application/xlsx" NC:saveToDisk="true" NC:alwaysAsk="false" ></RDF:Description>'''
    newcheckXls,oldcheckXls = checkContentInMimeTypes(checkThing='xls',mimetypepath=mimetypepath,addSet=addSet)
    #print('newcheckXls:%s'%(newcheckXls))
    #print('oldcheckXls:%s'%(oldcheckXls))
    newcheckXlsx,oldcheckXlsx = checkContentInMimeTypes(checkThing='xlsx',mimetypepath=mimetypepath,addSet=addSet)
    #print('newcheckXlsx:%s'%(newcheckXlsx))
    #print('oldcheckXlsx:%s'%(oldcheckXlsx))
def createDriver(requestInfo):
    '''create firefox webdriver and phantomjs webdriver 
        showMode:['show','hide']
    '''
    profile_directory = r'C:\Users\Administrator\AppData\Roaming\Mozilla\Firefox\Profiles\hdjophne.default'
    writeMimetype(mimetypepath=profile_directory)
    profile = webdriver.FirefoxProfile(profile_directory)
    profile.set_preference('browser.download.dir', r'E:\KM\GITPROJECT\HumanResourcesAss\AssForGaia\src\Guru\tmpfile')
    if requestInfo.showMode == 'show':    
        driver = webdriver.Firefox(profile)
    elif requestInfo.showMode == 'hide':
        driver = webdriver.PhantomJS('phantomjs')
    loginurl = requestInfo.rooturl + r'/Account/Logon?'
    driver.get(loginurl)
    time.sleep(2)
    driver.find_element_by_id('account_INPUT_IDENTITY').send_keys(requestInfo.webname)
    driver.find_element_by_id('password').send_keys(requestInfo.webpsw)
    driver.find_element_by_class_name('btn-lg').click()
    time.sleep(2)
    return driver
def readPublicByDoc(requestInfo,fileName):
    document = docx.Document(fileName)
    publicContent = []
    for table in document.tables:
        tableContent = {}
        codeValue = {}
        if table.cell(0,0).text == '公用代码配置':
            tableValue = []
            for row in table.rows:
                rowContent = []
                for cell in row.cells:
                    rowContent.append(cell.text)
                if rowContent[0] in ['公用代码配置','代码所属表','代码编号']:
                    pass
                elif rowContent[0] == '代码类型':
                    tableContent['codeType'] = rowContent[1]
                elif rowContent[0] == '代码名称':
                    tableContent['codeName'] = rowContent[1]
                else:
                    codeValue['codeItemID'] = rowContent[0]
                    codeValue['codeItemValue'] = rowContent[1]
                    tableValue.append(codeValue.copy())
            tableContent['codeItem'] = tableValue.copy()
            publicContent.append(tableContent.copy())
    #print('publicContent:%s'%(publicContent))
    return  publicContent     
def readExtendByDoc(requestInfo,fileName):
    document = docx.Document(fileName)
    extendContent = []
    for table in document.tables:
        tableContent = {}
        codeItem = []
        if table.cell(0,0).text == '业务数据配置':
            tableValue = []
            tableValueName = []
            for row in table.rows:
                rowContent = []
                for cell in row.cells:
                    rowContent.append(cell.text.replace('\u3000',''))
                if rowContent[1] in ['业务数据配置','项目']:
                    pass
                elif rowContent[1] == '业务数据类型':
                    tableContent['extendType'] = rowContent[2]
                elif rowContent[1] == '群组名称':
                    tableContent['extendName'] = rowContent[2]
                elif rowContent[1] == '重复限制':
                    tableContent['repeatLimit'] = rowContent[2]                    
                elif rowContent[1] == '群组描述':
                    tableContent['extendDesc'] = rowContent[2]
                elif rowContent[1] == '列字段名称':                                          
                    tableValueName = rowContent
                else:
                    tableValue.append(rowContent.copy())
            for singleValue in  tableValue:
                singleCodeItem = {}
                i = 0
                for i in range(len(singleValue)):
                    singleCodeItem[tableValueName[i]] = singleValue[i]   
                codeItem.append(singleCodeItem.copy())    
            tableContent['codeItem'] = codeItem.copy()
            extendContent.append(tableContent.copy())
    #print('extendContent:%s'%(extendContent))
    return  extendContent  
def queryPublicGroup(requestInfo):
    listPublicGroupName = []
    needurl = requestInfo.rooturl + r'/ePayroll/PayrollParameterSetting/PayrollPublicCodeMaintain.aspx'    
    content = http.request(uri = needurl,method = 'GET', headers=requestInfo.headers)[1]
    readSoup = bs.BeautifulSoup(content.decode(),'lxml')
    listPublicGroupName = listPublicGroupName + [singleGroup.find_all('span')[2].string for singleGroup in readSoup.find('div',id='M_UWebTreePublicCode_1').find_all('div') if len(singleGroup.find_all('span')[2].string)>0].copy()    
    needurl = requestInfo.rooturl + r'/ePayroll/PayrollParameterSetting/PayrollPublicCodeMaintain.aspx?Module=INS'     
    content = http.request(uri = needurl,method = 'GET', headers=requestInfo.headers)[1]
    readSoup = bs.BeautifulSoup(content.decode(),'lxml')
    listPublicGroupName = listPublicGroupName + [singleGroup.find_all('span')[2].string for singleGroup in readSoup.find('div',id='M_UWebTreePublicCode_1').find_all('div') if len(singleGroup.find_all('span')[2].string)>0].copy()    
    #print('listPublicGroupName:%s'%(listPublicGroupName))
    return listPublicGroupName
def addPublicGroup(driver,requestInfo,codeType,addPublicGroupName,listPublicGroupName):
    if addPublicGroupName in listPublicGroupName:
        print('Group name %s in PublicGroupName list,Add ignore!'%(addPublicGroupName))
    else:
        if codeType == '薪资公用代码':
            codeTypeValue ='PAY'
        elif codeType == '保险共用代码':
            codeTypeValue = 'INS'
        if  codeTypeValue is not None and len(addPublicGroupName) > 0:   
            needurl = requestInfo.rooturl +r'/ePayroll/PayrollParameterSetting/PayrollPublicCodeMaintainEdit.aspx?Module=' + codeTypeValue
            driver.get(needurl)
            time.sleep(2)
            driver.find_element_by_id('txtCode').send_keys(addPublicGroupName)
            time.sleep(2)
            driver.find_element_by_id('cmdAdd').click()
            time.sleep(2)
            message = driver.switch_to_alert().text
            if message == '新增成功！':
                print('Group Add Success:%s'%(addPublicGroupName))
                driver.switch_to_alert().accept()
            else:
                print('Group Add Failed:%s,Because of:%s'%(addPublicGroupName,message))
                driver.switch_to_alert().accept()
        else:
            print('codeTypeValue is None or addPublicGroupName:%s is null,Add ignore!'%(addPublicGroupName))
def updatePublicGroupByDoc(driver,requestInfo,listPublic):
    listPublicGroupName = queryPublicGroup(requestInfo)
    for singlePublic in listPublic:
        addPublicGroupCodeType = singlePublic.get('codeType')
        addPublicGroupName = singlePublic.get('codeName')
        if addPublicGroupCodeType in ['保险共用代码','薪资公用代码']:
            addPublicGroup(driver,requestInfo,addPublicGroupCodeType,addPublicGroupName,listPublicGroupName)
        else:
            print('for group:%s  Code Type is not 保险共用代码,薪资公用代码 ,add ignore'%(addPublicGroupName))
def queryExtendGroup(requestInfo):
    listExtendGroup = []
    needurl = requestInfo.rooturl + r'/ePayroll/PayrollParameterSetting/ExtendDataGroupSetting.aspx'    
    content = http.request(uri = needurl,method = 'GET', headers=requestInfo.headers)[1]
    readSoup = bs.BeautifulSoup(content.decode(),'lxml')
    singletr = readSoup.find('thead',class_='ig_e531842b_r1 WebGridText ig_e531842b_r4').find('tr')
    extendNamelist = [singletd.find('nobr').string for singletd in singletr.find_all('th')].copy()
    extendValuelist = []
    for singletr in readSoup.find('tbody',class_='ig_e531842b_r1 WebGridText ig_e531842b_r4').find_all('tr'):
        extendValuelist.append([singletd.find('nobr').string for singletd in singletr.find_all('td')].copy())
    if  len(extendValuelist) > 0: 
        for singleExtendValue in extendValuelist:
            singleExtendGroup = {}
            for i in range(len(singleExtendValue)):
                if extendNamelist[i] is not None and len(extendNamelist[i]) > 0 and extendNamelist[i] in ['群组名称','业务数据群组类型','群组描述','群组编号']:
                    singleExtendGroup[extendNamelist[i]] = singleExtendValue[i]
            listExtendGroup.append(singleExtendGroup.copy())
    else:
        print('No extend group list')      
    print('listExtendGroup:%s'%(listExtendGroup))
    return listExtendGroup
def queryExtendDetail(requestInfo,groupID):
    listExtendDetail = [] 
    needurl = requestInfo.rooturl + r'/ePayroll/PayrollParameterSetting/ExtendDataGroupSettingDetail.aspx?GroupID=' +  groupID   
    content = http.request(uri = needurl,method = 'GET', headers=requestInfo.headers)[1]
    readSoup = bs.BeautifulSoup(content.decode(),'lxml')
    singletr = readSoup.find('thead',class_='ig_e531842b_r1 WebGridText ig_e531842b_r4').find('tr')
    extendDetailNamelist = [singletd.find('nobr').string for singletd in singletr.find_all('th')].copy()
    extendDetailValuelist = []
    for singletr in readSoup.find('tbody',class_='ig_e531842b_r1 WebGridText ig_e531842b_r4').find_all('tr'):
        extendDetailValuelist.append([singletd.find('nobr').string for singletd in singletr.find_all('td')].copy())
    if  len(extendDetailValuelist) > 0: 
        for singleExtendDetailValue in extendDetailValuelist:
            singleExtendDetail = {}
            for i in range(len(singleExtendDetailValue)):
                if extendDetailNamelist[i] is not None and len(extendDetailNamelist[i]) > 0 and extendDetailNamelist[i] in ['序号','字段名称','数据类型','是否必填','显示方式','数据关联','字段编号']:
                    singleExtendDetail[extendDetailNamelist[i]] = singleExtendDetailValue[i].replace('\xa0','')
            listExtendDetail.append(singleExtendDetail.copy())
    else:
        print('No extend detail list')      
    print('listExtendDetail:%s'%(listExtendDetail))
    return listExtendDetail

              
def addExtendGroup(driver,requestInfo,addSingleExtendGroup,listExtendGroup):  
    addExtendGroupName = addSingleExtendGroup.get('extendName')
    listExtendGroupName = [group.get('群组名称') for group in listExtendGroup]
    if addExtendGroupName in listExtendGroupName:
        print('ExtendGroupName %s in ExtendGroupName list,Add ignore!'%(addExtendGroupName))
    elif len(addExtendGroupName) > 0:
        needurl = requestInfo.rooturl + r'/ePayroll/PayrollParameterSetting/ExtendDataGroupEdit.aspx'
        driver.get(needurl)
        time.sleep(2)
        driver.find_element_by_id('txtGroupName').send_keys(addExtendGroupName)
        addExtendGroupDesc = addSingleExtendGroup.get('extendDesc')
        driver.find_element_by_id('txtGroupDescription').send_keys(addExtendGroupDesc)
        addExtendGroupRepeatLimit = addSingleExtendGroup.get('repeatLimit')
        if addExtendGroupRepeatLimit in ['无限制','按日','按月','按季度','按年']:
            driver.find_element_by_xpath("//select[@id='DropDownListLimitType']/option[@title='"+ addExtendGroupRepeatLimit +"']").click()
        else:
            print('Group repeat limit:%s is not variable!'%(addExtendGroupRepeatLimit))
        addExtendGroupType = addSingleExtendGroup.get('extendType') 
        if addExtendGroupType == '组织业务数据':
            driver.find_element_by_id('rlType_0').click()
        elif addExtendGroupType == '员工业务数据':
            driver.find_element_by_id('rlType_1').click()
        else:
            print('Group type :%s is not variable!'%(addExtendGroupType))
        driver.find_element_by_id('btnAdd').click()
        time.sleep(2)
        message = driver.switch_to_alert().text
        print(message)
        if message == '新增成功！':
            print('Group Add Success:%s'%(addExtendGroupName))
            driver.switch_to_alert().accept()
        else:
            print('Group Add Failed:%s,Because of:%s'%(addExtendGroupName,message))
            driver.switch_to_alert().accept()
    else:
        print('addExtendGroupName:%s is null,Add ignore!'%(addExtendGroupName))
def addExtendDetail(driver,requestInfo,addSingleExtendDetail,listExtendDetail,groupID):        
    addExtendDetailName = addSingleExtendDetail.get('列字段名称')
    listExtendDetailName = [group.get('字段名称') for group in listExtendDetail]
    if addExtendDetailName in listExtendDetailName:
        print('ExtendDetailName %s in ExtendDetailName list,Add ignore!'%(addExtendDetailName))
    elif len(addExtendDetailName) > 0:
        needurl = requestInfo.rooturl + r'/ePayroll/PayrollParameterSetting/ExtendDataGroupDetailEdit.aspx?GroupID=' + groupID
        driver.get(needurl)
        time.sleep(2)
        driver.find_element_by_id('txtColumnName').send_keys(addExtendDetailName)
        
        addExtendDetailOrder = addSingleExtendDetail.get('序号')
        driver.find_element_by_id('txtOrder').send_keys(addExtendDetailOrder)
        
        addExtendDetailShowMode = addSingleExtendDetail.get('显示方式') 
        if addExtendDetailShowMode == '文本框':
            driver.find_element_by_id('rdoShowMode_0').click()
        elif addExtendDetailShowMode == '选择框':
            driver.find_element_by_id('rdoShowMode_1').click()
        else:
            print('GroupDetail show mode :%s is not variable!'%(addExtendDetailShowMode))     

        addExtendDetailMustFill = addSingleExtendDetail.get('是否必填') 
        if addExtendDetailMustFill == '是':
            driver.find_element_by_id('rdoMustFill_0').click()
        elif addExtendDetailMustFill == '否':
            driver.find_element_by_id('rdoMustFill_1').click()
        else:
            print('GroupDetail MustFill :%s is not variable!'%(addExtendDetailMustFill)) 

        addExtendDetailColumnType = addSingleExtendDetail.get('数据类型') 
        if addExtendDetailColumnType == '文本':
            driver.find_element_by_id('rdoColumnType_0').click()
        elif addExtendDetailColumnType == '日期':
            driver.find_element_by_id('rdoColumnType_1').click()
        elif addExtendDetailColumnType == '数字':
            driver.find_element_by_id('rdoColumnType_2').click()
        else:
            print('GroupDetail ColumnType :%s is not variable!'%(addExtendDetailColumnType)) 
            
        addExtendDetailPublicCode = addSingleExtendDetail.get('数据关联') 
        if  addExtendDetailShowMode == '选择框' and addExtendDetailPublicCode is not None and len(addExtendDetailPublicCode) > 0: 
            driver.find_element_by_xpath("//select[@id='drpPublicCode']/option[@title='"+ addExtendDetailPublicCode +"']").click()
    
        driver.find_element_by_id('btnAdd').click()
        time.sleep(2)
        message = driver.switch_to_alert().text
        print(message)
        if message == '新增成功！':
            print('GroupDetail Add Success:%s'%(addExtendDetailName))
            driver.switch_to_alert().accept()
        else:
            print('GroupDetail Add Failed:%s,Because of:%s'%(addExtendDetailName,message))
            driver.switch_to_alert().accept()
    else:
        print('addExtendDetailName:%s is null,Add ignore!'%(addExtendDetailName))        
           
            
 
    


cookie = cookielib.CookieJar()
handler = urllib.request.HTTPCookieProcessor(cookie)
CookieOpener = urllib.request.build_opener(handler)
headers = {}
headers['Accept'] = 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8'
headers['Accept-Encoding'] = 'deflate'
headers['Accept-Language'] = 'zh-CN,zh;q=0.8,en-US;q=0.5,en;q=0.3'
headers['Connection'] = 'keep-alive'
headers['User-Agent'] = 'Mozilla/5.0 (Windows NT 10.0; WOW64; rv:51.0) Gecko/20100101 Firefox/51.0'

rooturl = r'http://peter/zybxehr'
host = r'peter'
port = r'80'
webname = r'sa'
webpsw = r'111111'
tmpFilePath = r'E:\KM\GITPROJECT\HumanResourcesAss\AssForGaia\src\Guru\tmpfile'
showMode = r'show'
loginurl = rooturl + r'/Account/Logon?'
requestInfo = UrlRequest(rooturl=rooturl,host=host,port=port,webname=webname,webpsw=webpsw,tmpFilePath=tmpFilePath,showMode=showMode,cookie=cookie,headers=headers,loginurl=loginurl,CookieOpener=CookieOpener)
http = httplib2.Http()


resLogin = loginInSystem(requestInfo)
driver = createDriver(requestInfo)
#事项1：通过doc读取公用代码
fileName = r'E:\工作文档\1608-中银\06业务流程分析\薪资管理\test.docx'
#listPublicByDoc = readPublicByDoc(requestInfo,fileName)
#listExtendByDoc = readExtendByDoc(requestInfo,fileName)
#事项2：读取已存在的保险公用代码及薪资公用代码 更新公用代码群组时会自动调用
#listPublicGroupName = queryPublicGroup(requestInfo)
#事项3：新增公用代码群组 更新公用代码群组时会自动调用
#resAddPublicGroup = addPublicGroup(driver,requestInfo,'薪资公用代码','test002')
#事项4：更新公用代码群组
#updatePublicGroupByDoc(driver,requestInfo,listPublicByDoc)
addSingleExtendGroup = {'codeItem': [{'列字段名称': '人员险种', '数据关联': '人员险种(人事同步)', '是否必填': '是', '数据类型': '文本', '显示方式': '选择框', '序号': '1'}, {'列字段名称': '财务代码', '数据关联': '', '是否必填': '是', '数据类型': '文本', '显示方式': '文本框', '序号': '2'}], 'extendDesc': '为费用凭证报表产品代码列维护员工险种编码与财务码对应关系，维护在根组织（0总公司）上。', 'extendType': '组织业务数据', 'extendName': '员工险种-财务码', 'repeatLimit': '无限制'}
addSingleExtendDetail = addSingleExtendGroup.get('codeItem')[0]
#addExtendGroup(driver,requestInfo,addSingleExtendGroup,listExtendGroupName)
#queryExtendGroup(requestInfo)
listExtendDetail = queryExtendDetail(requestInfo,'90184dd3-25ba-4475-a15a-ab5e0811071d')

addExtendDetail(driver,requestInfo,addSingleExtendDetail,listExtendDetail,'90184dd3-25ba-4475-a15a-ab5e0811071d')

t2 = time.clock()
print ("the project costs %.9fs"%(t2-t1))